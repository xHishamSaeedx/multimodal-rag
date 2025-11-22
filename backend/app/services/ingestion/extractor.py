"""
Text extraction service.

Extracts text from PDF, DOCX, TXT, MD files.
"""

import io
from pathlib import Path
from typing import BinaryIO
from dataclasses import dataclass
from datetime import datetime

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from app.utils.exceptions import (
    ExtractionError,
    UnsupportedFileTypeError,
    FileReadError,
)


@dataclass
class ExtractedContent:
    """Represents extracted content from a document with metadata."""
    
    text: str
    file_name: str
    file_type: str
    file_size: int | None = None
    page_count: int | None = None
    extracted_at: datetime | None = None
    metadata: dict | None = None
    
    def __post_init__(self):
        """Set default values."""
        if self.extracted_at is None:
            self.extracted_at = datetime.utcnow()
        if self.metadata is None:
            self.metadata = {}


class TextExtractor:
    """
    Service for extracting text from various document formats.
    
    Supports:
    - PDF: Using PyMuPDF (fitz)
    - DOCX: Using python-docx
    - TXT/MD: Direct file read
    """
    
    # Supported MIME types and file extensions
    SUPPORTED_MIME_TYPES = {
        "application/pdf": [".pdf"],
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"],
        "text/plain": [".txt"],
        "text/markdown": [".md", ".markdown"],
    }
    
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".txt": "txt",
        ".md": "md",
        ".markdown": "md",
    }
    
    def __init__(self):
        """Initialize the text extractor."""
        if fitz is None:
            raise ImportError(
                "PyMuPDF is not installed. Install it with: pip install PyMuPDF"
            )
        if DocxDocument is None:
            raise ImportError(
                "python-docx is not installed. Install it with: pip install python-docx"
            )
    
    def extract_from_file(
        self,
        file_path: str | Path,
        file_type: str | None = None,
    ) -> ExtractedContent:
        """
        Extract text from a file path.
        
        Args:
            file_path: Path to the file to extract text from
            file_type: Optional file type hint (pdf, docx, txt, md)
                      If not provided, will be inferred from file extension
        
        Returns:
            ExtractedContent object with text and metadata
        
        Raises:
            UnsupportedFileTypeError: If file type is not supported
            FileReadError: If file cannot be read
            ExtractionError: If extraction fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileReadError(
                f"File not found: {file_path}",
                {"file_path": str(file_path)},
            )
        
        # Determine file type
        if file_type is None:
            file_type = self._infer_file_type(file_path)
        
        if file_type not in self.SUPPORTED_EXTENSIONS.values():
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS.values())}",
                {"file_path": str(file_path), "file_type": file_type},
            )
        
        # Get file size
        file_size = file_path.stat().st_size
        
        # Extract based on file type
        try:
            if file_type == "pdf":
                text, page_count, metadata = self._extract_pdf(file_path)
            elif file_type == "docx":
                text, page_count, metadata = self._extract_docx(file_path)
            elif file_type in ("txt", "md"):
                text, page_count, metadata = self._extract_text(file_path)
            else:
                raise UnsupportedFileTypeError(
                    f"Unsupported file type: {file_type}",
                    {"file_path": str(file_path), "file_type": file_type},
                )
        except Exception as e:
            if isinstance(e, (UnsupportedFileTypeError, FileReadError, ExtractionError)):
                raise
            raise ExtractionError(
                f"Failed to extract text from {file_path}: {str(e)}",
                {"file_path": str(file_path), "file_type": file_type, "error": str(e)},
            ) from e
        
        return ExtractedContent(
            text=text,
            file_name=file_path.name,
            file_type=file_type,
            file_size=file_size,
            page_count=page_count,
            metadata=metadata,
        )
    
    def extract_from_bytes(
        self,
        file_bytes: bytes,
        file_name: str,
        file_type: str | None = None,
    ) -> ExtractedContent:
        """
        Extract text from file bytes.
        
        Args:
            file_bytes: File content as bytes
            file_name: Name of the file (used to infer type)
            file_type: Optional file type hint (pdf, docx, txt, md)
                      If not provided, will be inferred from file name
        
        Returns:
            ExtractedContent object with text and metadata
        
        Raises:
            UnsupportedFileTypeError: If file type is not supported
            ExtractionError: If extraction fails
        """
        # Determine file type
        if file_type is None:
            file_path = Path(file_name)
            file_type = self._infer_file_type(file_path)
        
        if file_type not in self.SUPPORTED_EXTENSIONS.values():
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS.values())}",
                {"file_name": file_name, "file_type": file_type},
            )
        
        file_size = len(file_bytes)
        
        # Extract based on file type
        try:
            file_io = io.BytesIO(file_bytes)
            
            if file_type == "pdf":
                text, page_count, metadata = self._extract_pdf(file_io)
            elif file_type == "docx":
                text, page_count, metadata = self._extract_docx(file_io)
            elif file_type in ("txt", "md"):
                text, page_count, metadata = self._extract_text(file_io)
            else:
                raise UnsupportedFileTypeError(
                    f"Unsupported file type: {file_type}",
                    {"file_name": file_name, "file_type": file_type},
                )
        except Exception as e:
            if isinstance(e, (UnsupportedFileTypeError, ExtractionError)):
                raise
            raise ExtractionError(
                f"Failed to extract text from {file_name}: {str(e)}",
                {"file_name": file_name, "file_type": file_type, "error": str(e)},
            ) from e
        
        return ExtractedContent(
            text=text,
            file_name=file_name,
            file_type=file_type,
            file_size=file_size,
            page_count=page_count,
            metadata=metadata,
        )
    
    def _infer_file_type(self, file_path: Path) -> str:
        """
        Infer file type from file extension.
        
        Args:
            file_path: Path to the file
        
        Returns:
            File type string (pdf, docx, txt, md)
        """
        ext = file_path.suffix.lower()
        file_type = self.SUPPORTED_EXTENSIONS.get(ext)
        
        if file_type is None:
            raise UnsupportedFileTypeError(
                f"Unsupported file extension: {ext}. "
                f"Supported extensions: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}",
                {"file_path": str(file_path), "extension": ext},
            )
        
        return file_type
    
    def _extract_pdf(
        self, source: Path | BinaryIO
    ) -> tuple[str, int, dict]:
        """
        Extract text from a PDF file using PyMuPDF (fitz).
        
        Args:
            source: Path to PDF file or BytesIO object
        
        Returns:
            Tuple of (text, page_count, metadata)
        """
        try:
            # Handle BytesIO or Path objects
            if isinstance(source, Path):
                # Path object - pass as string
                doc = fitz.open(str(source))
            elif hasattr(source, 'read'):
                # File-like object (BytesIO, etc.) - read bytes and pass to fitz.open
                source.seek(0)  # Reset to beginning in case it was read before
                pdf_bytes = source.read()
                # PyMuPDF can accept bytes directly or via stream parameter
                try:
                    # Try direct bytes first (newer PyMuPDF versions)
                    doc = fitz.open(pdf_bytes)
                except (TypeError, ValueError):
                    # Fall back to stream parameter (older versions)
                    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            else:
                # Assume it's bytes already
                doc = fitz.open(source)
            text_parts = []
            metadata = {}
            
            # Extract document metadata if available
            if doc.metadata:
                metadata.update({
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "creator": doc.metadata.get("creator", ""),
                    "producer": doc.metadata.get("producer", ""),
                    "creation_date": doc.metadata.get("creationDate", ""),
                    "modification_date": doc.metadata.get("modDate", ""),
                })
            
            # Get page count before closing
            page_count = len(doc)
            
            # Extract text from each page
            for page_num in range(page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Add page separator for multi-page documents
                if page_text.strip():
                    if text_parts:  # Not the first page
                        text_parts.append("\n\n--- Page Break ---\n\n")
                    text_parts.append(page_text)
            
            text = "".join(text_parts)
            
            doc.close()
            
            # Clean up text
            text = self._clean_text(text)
            
            return text, page_count, metadata
        
        except Exception as e:
            raise ExtractionError(
                f"Failed to extract text from PDF: {str(e)}",
                {"source": str(source), "error": str(e)},
            ) from e
    
    def _extract_docx(
        self, source: Path | BinaryIO
    ) -> tuple[str, int, dict]:
        """
        Extract text from a DOCX file using python-docx.
        
        Args:
            source: Path to DOCX file or BytesIO object
        
        Returns:
            Tuple of (text, page_count, metadata)
        """
        try:
            doc = DocxDocument(source)
            text_parts = []
            metadata = {}
            
            # Extract document metadata
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "last_modified_by": core_props.last_modified_by or "",
                })
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            
            # Clean up text
            text = self._clean_text(text)
            
            # DOCX doesn't have explicit page count, estimate based on content
            page_count = None
            
            return text, page_count, metadata
        
        except Exception as e:
            raise ExtractionError(
                f"Failed to extract text from DOCX: {str(e)}",
                {"source": str(source), "error": str(e)},
            ) from e
    
    def _extract_text(
        self, source: Path | BinaryIO
    ) -> tuple[str, int, dict]:
        """
        Extract text from a TXT or MD file by direct read.
        
        Args:
            source: Path to text file or BytesIO object
        
        Returns:
            Tuple of (text, page_count, metadata)
        """
        try:
            if isinstance(source, Path):
                # Read from file path
                with open(source, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                # Read from BytesIO
                text = source.read().decode("utf-8")
            
            # Clean up text
            text = self._clean_text(text)
            
            # Simple text files don't have page count
            page_count = None
            metadata = {}
            
            return text, page_count, metadata
        
        except UnicodeDecodeError:
            # Try with different encodings
            try:
                if isinstance(source, Path):
                    with open(source, "r", encoding="latin-1") as f:
                        text = f.read()
                else:
                    source.seek(0)
                    text = source.read().decode("latin-1")
                
                text = self._clean_text(text)
                return text, None, {}
            except Exception as e:
                raise FileReadError(
                    f"Failed to read text file with supported encodings: {str(e)}",
                    {"source": str(source), "error": str(e)},
                ) from e
        except Exception as e:
            raise FileReadError(
                f"Failed to read text file: {str(e)}",
                {"source": str(source), "error": str(e)},
            ) from e
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by normalizing whitespace.
        
        Args:
            text: Raw extracted text
        
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        
        # Remove excessive blank lines (more than 2 consecutive)
        import re
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def is_supported(self, file_path: str | Path) -> bool:
        """
        Check if a file type is supported for extraction.
        
        Args:
            file_path: Path to the file
        
        Returns:
            True if file type is supported, False otherwise
        """
        try:
            file_path = Path(file_path)
            ext = file_path.suffix.lower()
            return ext in self.SUPPORTED_EXTENSIONS
        except Exception:
            return False
